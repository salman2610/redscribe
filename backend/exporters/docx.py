from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from datetime import datetime
import uuid

SEVERITY_COLORS = {
    'critical': (239, 68, 68),      # #ef4444
    'high': (249, 115, 22),          # #f97316
    'medium': (234, 179, 8),         # #eab308
    'low': (59, 130, 246),           # #3b82f6
    'info': (107, 114, 128),         # #6b7280
}

SEVERITY_BG = {
    'critical': (254, 242, 242),     # #fef2f2
    'high': (255, 247, 237),         # #fff7ed
    'medium': (254, 252, 232),       # #fefce8
    'low': (239, 246, 255),          # #eff6ff
    'info': (249, 250, 251),         # #f9fafb
}

RED_SCRIBE_RED = RGBColor(187, 38, 73)  # #BB2649
DARK_BG = RGBColor(8, 8, 8)            # #080808
LIGHT_TEXT = RGBColor(17, 17, 17)      # #111

SEV_ORDER = ['critical', 'high', 'medium', 'low', 'info']


def get_sev_index(sev):
    try:
        return SEV_ORDER.index(sev.lower() if sev else 'info')
    except ValueError:
        return 99


def shade_cell(cell, fill_color):
    """Add background color to table cell"""
    tcPr = cell._element.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:shd')
    tcVAlign.set(qn('w:fill'), fill_color)
    tcPr.append(tcVAlign)


def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()


def add_cover_page(doc, project, finding_counts):
    """Add professional cover page"""
    # Dark background section (simulated with table)
    cover_table = doc.add_table(rows=5, cols=1)
    cover_table.autofit = False
    cover_table.allow_autofit = False
    
    # Set width to full page
    tbl = cover_table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Logo row
    logo_cell = cover_table.rows[0].cells[0]
    shade_cell(logo_cell, "080808")
    logo_para = logo_cell.paragraphs[0]
    logo_run = logo_para.add_run("Red")
    logo_run.font.size = Pt(24)
    logo_run.font.bold = True
    logo_run.font.color.rgb = RGBColor(255, 255, 255)
    logo_run2 = logo_para.add_run("Scribe")
    logo_run2.font.size = Pt(24)
    logo_run2.font.bold = True
    logo_run2.font.color.rgb = RED_SCRIBE_RED
    logo_cell.paragraphs[0].paragraph_format.space_before = Pt(12)
    logo_cell.paragraphs[0].paragraph_format.space_after = Pt(12)
    
    # Title section
    title_cell = cover_table.rows[1].cells[0]
    shade_cell(title_cell, "080808")
    title_para = title_cell.paragraphs[0]
    
    # Label
    label_run = title_para.add_run("// PENETRATION TEST REPORT\n")
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = RED_SCRIBE_RED
    
    # Project name
    name_run = title_para.add_run(f"{project.get('name', 'Security Assessment')}\n")
    name_run.font.size = Pt(36)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Client
    client_run = title_para.add_run(f"{project.get('client_name', '')}\n")
    client_run.font.size = Pt(12)
    client_run.font.color.rgb = RGBColor(100, 100, 100)
    
    title_cell.paragraphs[0].paragraph_format.space_before = Pt(20)
    title_cell.paragraphs[0].paragraph_format.space_after = Pt(20)
    
    # Stats row
    stats_cell = cover_table.rows[2].cells[0]
    shade_cell(stats_cell, "080808")
    stats_para = stats_cell.paragraphs[0]
    
    total = sum(finding_counts.values())
    stats_text = f"Total: {total}  |  Critical: {finding_counts['critical']}  |  High: {finding_counts['high']}  |  Medium: {finding_counts['medium']}  |  Low: {finding_counts['low']}"
    stats_run = stats_para.add_run(stats_text)
    stats_run.font.size = Pt(12)
    stats_run.font.bold = True
    stats_run.font.color.rgb = RGBColor(255, 255, 255)
    stats_cell.paragraphs[0].paragraph_format.space_before = Pt(20)
    stats_cell.paragraphs[0].paragraph_format.space_after = Pt(20)
    
    # Footer row
    footer_cell = cover_table.rows[3].cells[0]
    shade_cell(footer_cell, "080808")
    footer_para = footer_cell.paragraphs[0]
    
    date_run = footer_para.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}\n")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(80, 80, 80)
    
    scope_run = footer_para.add_run(f"Scope: {project.get('scope', 'N/A')}")
    scope_run.font.size = Pt(10)
    scope_run.font.color.rgb = RGBColor(80, 80, 80)
    
    footer_cell.paragraphs[0].paragraph_format.space_before = Pt(20)
    footer_cell.paragraphs[0].paragraph_format.space_after = Pt(12)
    
    # Remove table borders (visual background effect)
    for row in cover_table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    add_page_break(doc)


def add_engagement_page(doc, project, findings, executive_summary=''):
    """Add engagement details and risk overview"""
    # Header
    header_para = doc.add_paragraph()
    run = header_para.add_run("Red")
    run.font.size = Pt(14)
    run.font.bold = True
    run = header_para.add_run("Scribe")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RED_SCRIBE_RED
    header_para.paragraph_format.space_after = Pt(6)
    
    # Engagement Details
    title = doc.add_paragraph("// ENGAGEMENT DETAILS")
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(12)
    title_run = title.runs[0]
    title_run.font.size = Pt(10)
    title_run.font.bold = True
    title_run.font.color.rgb = RED_SCRIBE_RED
    
    # Details table
    details_table = doc.add_table(rows=5, cols=2)
    details_table.style = 'Light Grid'
    details_table.autofit = False
    
    details = [
        ("Project", project.get("name", "N/A")),
        ("Client", project.get("client_name", "N/A")),
        ("Scope", project.get("scope", "N/A")),
        ("Report Date", datetime.now().strftime("%d %B %Y")),
        ("Total Findings", str(len(findings))),
    ]
    
    for i, (label, value) in enumerate(details):
        label_cell = details_table.rows[i].cells[0]
        value_cell = details_table.rows[i].cells[1]
        
        label_para = label_cell.paragraphs[0]
        label_para.add_run(label)
        label_para.runs[0].font.size = Pt(10)
        label_para.runs[0].font.bold = True
        label_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
        
        value_para = value_cell.paragraphs[0]
        value_para.add_run(value)
        value_para.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Risk Overview
    risk_title = doc.add_paragraph("// RISK OVERVIEW")
    risk_title.paragraph_format.space_before = Pt(12)
    risk_title.paragraph_format.space_after = Pt(12)
    risk_title_run = risk_title.runs[0]
    risk_title_run.font.size = Pt(10)
    risk_title_run.font.bold = True
    risk_title_run.font.color.rgb = RED_SCRIBE_RED
    
    # Risk cards as table
    counts = {s: sum(1 for f in findings if (f.get('severity') or 'info').lower() == s) for s in SEV_ORDER}
    risk_table = doc.add_table(rows=1, cols=5)
    risk_table.autofit = False
    
    for i, sev in enumerate(SEV_ORDER):
        cell = risk_table.rows[0].cells[i]
        shade_cell(cell, "f9fafb")
        
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        count_run = para.add_run(str(counts[sev]))
        count_run.font.size = Pt(20)
        count_run.font.bold = True
        color = SEVERITY_COLORS[sev]
        count_run.font.color.rgb = RGBColor(*color)
        
        para.add_run("\n")
        label_run = para.add_run(sev.upper())
        label_run.font.size = Pt(9)
        label_run.font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Executive Summary
    if executive_summary:
        summary_title = doc.add_paragraph("// EXECUTIVE SUMMARY")
        summary_title.paragraph_format.space_before = Pt(12)
        summary_title.paragraph_format.space_after = Pt(12)
        summary_title_run = summary_title.runs[0]
        summary_title_run.font.size = Pt(10)
        summary_title_run.font.bold = True
        summary_title_run.font.color.rgb = RED_SCRIBE_RED
        
        summary_para = doc.add_paragraph(executive_summary)
        summary_para.paragraph_format.space_after = Pt(12)
        summary_para.runs[0].font.size = Pt(11)
    
    add_page_break(doc)


def add_findings_page(doc, findings):
    """Add all findings"""
    # Header
    header_para = doc.add_paragraph()
    run = header_para.add_run("Red")
    run.font.size = Pt(14)
    run.font.bold = True
    run = header_para.add_run("Scribe")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RED_SCRIBE_RED
    header_para.paragraph_format.space_after = Pt(6)
    
    # Findings title
    title = doc.add_paragraph(f"// FINDINGS ({len(findings)} total)")
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(12)
    title_run = title.runs[0]
    title_run.font.size = Pt(10)
    title_run.font.bold = True
    title_run.font.color.rgb = RED_SCRIBE_RED
    
    # Sort findings by severity
    sorted_findings = sorted(findings, key=lambda x: get_sev_index(x.get('severity', 'info')))
    
    for i, f in enumerate(sorted_findings):
        sev = (f.get('severity') or 'info').lower()
        
        # Finding header
        header = doc.add_paragraph()
        header.paragraph_format.space_before = Pt(12)
        header.paragraph_format.space_after = Pt(6)
        
        # Number + Title
        num_run = header.add_run(f"F{str(i+1).zfill(2)}: ")
        num_run.font.size = Pt(11)
        num_run.font.bold = True
        num_run.font.color.rgb = RGBColor(150, 150, 150)
        
        title_run = header.add_run(f"{f.get('title', 'Untitled')}")
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        
        # Severity badge
        severity_run = header.add_run(f"  [{sev.upper()}]")
        severity_run.font.size = Pt(10)
        severity_run.font.bold = True
        color = SEVERITY_COLORS.get(sev, (107, 114, 128))
        severity_run.font.color.rgb = RGBColor(*color)
        
        # Metadata
        meta_parts = []
        if f.get('cvss_score'):
            meta_parts.append(f"CVSS: {f['cvss_score']}")
        if f.get('cwe'):
            meta_parts.append(f"{f['cwe']}")
        if f.get('owasp'):
            meta_parts.append(f"{f['owasp']}")
        if f.get('source_tool'):
            meta_parts.append(f"via {f['source_tool']}")
        
        if meta_parts:
            meta_para = doc.add_paragraph(" | ".join(meta_parts))
            meta_para.paragraph_format.space_before = Pt(0)
            meta_para.paragraph_format.space_after = Pt(6)
            meta_para.runs[0].font.size = Pt(9)
            meta_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
        
        # Description
        if f.get('description'):
            desc_label = doc.add_paragraph()
            desc_label.paragraph_format.space_before = Pt(6)
            desc_label.paragraph_format.space_after = Pt(3)
            desc_run = desc_label.add_run("Description")
            desc_run.font.size = Pt(9)
            desc_run.font.bold = True
            desc_run.font.color.rgb = RGBColor(100, 100, 100)
            
            desc_text = doc.add_paragraph(f.get('description', ''))
            desc_text.paragraph_format.space_before = Pt(0)
            desc_text.paragraph_format.space_after = Pt(6)
            desc_text.paragraph_format.left_indent = Inches(0.25)
            desc_text.runs[0].font.size = Pt(10)
        
        # Business Impact
        if f.get('business_impact'):
            impact_label = doc.add_paragraph()
            impact_label.paragraph_format.space_before = Pt(6)
            impact_label.paragraph_format.space_after = Pt(3)
            impact_run = impact_label.add_run("Business Impact")
            impact_run.font.size = Pt(9)
            impact_run.font.bold = True
            impact_run.font.color.rgb = RGBColor(124, 58, 237)  # Purple
            
            impact_text = doc.add_paragraph(f.get('business_impact', ''))
            impact_text.paragraph_format.space_before = Pt(0)
            impact_text.paragraph_format.space_after = Pt(6)
            impact_text.paragraph_format.left_indent = Inches(0.25)
            impact_text.runs[0].font.size = Pt(10)
        
        # Remediation
        if f.get('remediation'):
            rem_label = doc.add_paragraph()
            rem_label.paragraph_format.space_before = Pt(6)
            rem_label.paragraph_format.space_after = Pt(3)
            rem_run = rem_label.add_run("Remediation")
            rem_run.font.size = Pt(9)
            rem_run.font.bold = True
            rem_run.font.color.rgb = RGBColor(5, 150, 105)  # Green
            
            rem_text = doc.add_paragraph(f.get('remediation', ''))
            rem_text.paragraph_format.space_before = Pt(0)
            rem_text.paragraph_format.space_after = Pt(6)
            rem_text.paragraph_format.left_indent = Inches(0.25)
            rem_text.runs[0].font.size = Pt(10)
        
        # Attack Scenario
        if f.get('attack_scenario'):
            attack_label = doc.add_paragraph()
            attack_label.paragraph_format.space_before = Pt(6)
            attack_label.paragraph_format.space_after = Pt(3)
            attack_run = attack_label.add_run("Attack Scenario")
            attack_run.font.size = Pt(9)
            attack_run.font.bold = True
            attack_run.font.color.rgb = RGBColor(220, 38, 38)  # Red
            
            attack_text = doc.add_paragraph(f.get('attack_scenario', ''))
            attack_text.paragraph_format.space_before = Pt(0)
            attack_text.paragraph_format.space_after = Pt(6)
            attack_text.paragraph_format.left_indent = Inches(0.25)
            attack_text.runs[0].font.size = Pt(10)
        
        # Affected hosts/ports tags
        tags = []
        if f.get('affected_hosts'):
            tags.append(", ".join(f.get('affected_hosts', [])))
        if f.get('affected_ports'):
            tags.append(f"Port {', '.join(str(p) for p in f.get('affected_ports', []))}")
        if f.get('cvss_vector'):
            tags.append(f.get('cvss_vector'))
        
        if tags:
            tags_para = doc.add_paragraph(" | ".join(tags))
            tags_para.paragraph_format.space_before = Pt(6)
            tags_para.paragraph_format.space_after = Pt(12)
            tags_para.runs[0].font.size = Pt(9)
            tags_para.runs[0].font.italic = True
            tags_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)


async def generate_docx(project_id: str, db: AsyncSession, executive_summary: str = '') -> bytes:
    """Generate DOCX report"""
    from models.models import Project, Finding

    proj_result = await db.execute(select(Project).where(Project.id == uuid.UUID(project_id)))
    project = proj_result.scalar_one_or_none()

    find_result = await db.execute(select(Finding).where(Finding.project_id == uuid.UUID(project_id)))
    findings = find_result.scalars().all()

    project_dict = {
        "name": project.name,
        "client_name": project.client_name,
        "scope": project.scope,
        "ai_provider": project.ai_provider,
    }

    findings_list = [{
        "title": f.title,
        "severity": f.severity.value.lower() if f.severity else "info",
        "cvss_score": f.cvss_score,
        "cvss_vector": f.cvss_vector,
        "cwe": f.cwe,
        "owasp": f.owasp,
        "description": f.description,
        "business_impact": f.business_impact,
        "remediation": f.remediation,
        "attack_scenario": f.attack_scenario,
        "affected_hosts": f.affected_hosts,
        "affected_ports": f.affected_ports,
        "source_tool": f.source_tool,
    } for f in findings]

    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Calculate finding counts
    counts = {s: sum(1 for f in findings_list if (f.get('severity') or 'info').lower() == s) for s in SEV_ORDER}
    
    # Build pages
    add_cover_page(doc, project_dict, counts)
    add_engagement_page(doc, project_dict, findings_list, executive_summary)
    add_findings_page(doc, findings_list)
    
    # Save to bytes
    from io import BytesIO
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    return docx_bytes.getvalue()
